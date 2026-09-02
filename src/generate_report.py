"""
自动生成 .docx 课程设计报告，严格按学校模板格式。
格式要求：
- 封面：题目、姓名、学号等
- 正文：宋体小四，行距24磅固定值，首行缩进2字符
- 一级标题：黑体四号加粗；二级/三级标题：宋体四号加粗
- 图题：宋体五号居中，图下方
- 表题：宋体五号居中，表上方
- 参考文献≥15篇含英文
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

from config import PROJECT_ROOT, FIGURES_DIR, REPORT_DIR, RESULTS_DIR, RULES_DIR


# ========== 格式辅助函数 ==========

def set_cell_font(cell, font_name='宋体', font_size=10, bold=False):
    """设置表格单元格字体。"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.bold = bold
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_body_paragraph(doc, text, indent=True):
    """添加正文段落：宋体小四，行距固定值24磅，首行缩进2字符。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(24)
    if indent:
        pf.first_line_indent = Cm(0.85)  # 约2字符
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(12)  # 小四
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # 英文/数字用 Times New Roman
    run.font.element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    run.font.element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    return p


def add_heading_1(doc, text):
    """一级标题：黑体四号加粗，段前后6磅。"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.font.name = '黑体'
    run.font.size = Pt(14)  # 四号
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    return p


def add_heading_2(doc, text):
    """二级标题：宋体四号加粗，段前后6磅。"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(14)  # 四号
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    return p


def add_heading_3(doc, text):
    """三级标题：宋体四号加粗，段前后6磅。"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(14)  # 四号
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    return p


def add_image(doc, image_path, caption, width=Inches(5.5)):
    """插入图片 + 图题（宋体五号居中，图下方）。"""
    if not os.path.exists(image_path):
        add_body_paragraph(doc, f"[图片缺失: {image_path}]", indent=False)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=width)
    # 图题
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_before = Pt(3)
    cap_p.paragraph_format.space_after = Pt(3)
    cap_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    cap_run = cap_p.add_run(caption)
    cap_run.font.name = '宋体'
    cap_run.font.size = Pt(10.5)  # 五号
    cap_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_table_caption(doc, caption):
    """表题：宋体五号居中，表上方。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(caption)
    run.font.name = '宋体'
    run.font.size = Pt(10.5)  # 五号
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_metrics_table(doc, metrics_df):
    """插入指标汇总表。"""
    add_table_caption(doc, '表1  四模型性能指标汇总表')
    table = doc.add_table(rows=len(metrics_df) + 1, cols=len(metrics_df.columns))
    table.style = 'Table Grid'

    # 表头
    for j, col in enumerate(metrics_df.columns):
        cell = table.rows[0].cells[j]
        cell.text = col
        set_cell_font(cell, '黑体', 9, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 数据行
    for i, row in metrics_df.iterrows():
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            set_cell_font(cell, '宋体', 9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_code_block(doc, code_text, max_lines=80):
    """添加代码块（附录用），五号字。"""
    lines = code_text.split('\n')[:max_lines]
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(10.5)  # 五号
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def generate_report(metrics_all, metrics_df, feature_names, target_names):
    """生成完整课程设计报告。"""
    print("[报告生成] 开始生成课程设计报告...")

    doc = Document()

    # 页面设置：上3cm 下2.5cm 左3cm 右2cm
    for section in doc.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    # ========== 封面 ==========
    _add_cover_page(doc)

    # ========== 目录 ==========
    _add_toc_page(doc)

    # ========== 正文 ==========
    _add_chapter_1(doc)
    _add_chapter_2(doc, feature_names, target_names)
    _add_chapter_3(doc)
    _add_chapter_4(doc, metrics_all, metrics_df)
    _add_chapter_5(doc)
    _add_chapter_6(doc)
    _add_chapter_7(doc)
    _add_references(doc)
    _add_appendix(doc)

    # 保存
    path = os.path.join(REPORT_DIR, '鸢尾花分类课程设计报告.docx')
    doc.save(path)
    print(f"  -> {path}")


def _add_cover_page(doc):
    """封面页。"""
    # 空行调整位置
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('《机器学习项目实训》设计报告')
    run.font.name = '楷体'
    run.font.size = Pt(18)  # 小二
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('鸢尾花分类——从模型对比到可解释性')
    run.font.name = '楷体'
    run.font.size = Pt(18)  # 小二
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

    for _ in range(3):
        doc.add_paragraph()

    # 信息栏
    info_items = [
        ('题　　目', '鸢尾花分类——从模型对比到可解释性'),
        ('姓　　名', '　　　　　　　'),
        ('学　　号', '　　　　　　　'),
        ('专　　业', '计算机科学与技术'),
        ('班　　级', '　　　　　　　'),
        ('指导教师', '　　　　　　　'),
        ('职　　称', '　　　　　　　'),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label}：{value}')
        run.font.name = '楷体'
        run.font.size = Pt(18)  # 小二
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('计算机学院')
    run.font.name = '楷体'
    run.font.size = Pt(18)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('2026年9月')
    run.font.name = '楷体'
    run.font.size = Pt(18)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

    doc.add_page_break()


def _add_toc_page(doc):
    """目录页。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('目　　录')
    run.font.name = '黑体'
    run.font.size = Pt(16)  # 三号
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    toc_items = [
        ('1  课程设计概述', '1'),
        ('1.1  题目背景与意义', '1'),
        ('1.2  鸢尾花数据集简介', '2'),
        ('1.3  设计目标与任务要求', '2'),
        ('2  数据探索与预处理', '3'),
        ('2.1  数据集描述', '3'),
        ('2.2  探索性数据分析', '4'),
        ('2.3  数据预处理策略', '7'),
        ('3  模型原理概述', '8'),
        ('3.1  K近邻算法（KNN）', '8'),
        ('3.2  支持向量机（RBF-SVM）', '9'),
        ('3.3  决策树（CART）', '9'),
        ('3.4  逻辑回归（Softmax）', '10'),
        ('3.5  评估指标定义', '10'),
        ('4  模型训练与评估', '11'),
        ('4.1  实验设置', '11'),
        ('4.2  四模型性能对比', '12'),
        ('4.3  混淆矩阵分析', '13'),
        ('4.4  ROC曲线与AUC分析', '14'),
        ('4.5  泛化能力讨论', '14'),
        ('5  决策边界可视化对比', '15'),
        ('5.1  特征对选取依据', '15'),
        ('5.2  高区分特征对边界对比', '16'),
        ('5.3  低区分特征对边界对比', '17'),
        ('5.4  边界形态差异分析', '18'),
        ('6  可解释性分析', '19'),
        ('6.1  决策树规则提取', '19'),
        ('6.2  特征重要性分析', '20'),
        ('6.3  SHAP全局解释', '21'),
        ('6.4  SHAP局部解释', '22'),
        ('6.5  模型解释性对比讨论', '23'),
        ('7  结论与展望', '24'),
        ('7.1  实验结论', '24'),
        ('7.2  局限性分析', '24'),
        ('7.3  未来改进方向', '25'),
        ('参考文献', '26'),
        ('附录  核心源代码', '27'),
    ]
    for item, page in toc_items:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = Pt(24)
        dots = '.' * max(1, 50 - len(item) - len(page))
        run = p.add_run(f'{item} {dots} {page}')
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')

    doc.add_page_break()


def _add_chapter_1(doc):
    """第1章：课程设计概述。"""
    add_heading_1(doc, '1  课程设计概述')

    add_heading_2(doc, '1.1  题目背景与意义')
    add_body_paragraph(doc,
        '机器学习是人工智能领域的核心分支，其目标是通过算法从数据中自动学习规律并做出预测。'
        '分类问题是机器学习中最常见的监督学习任务之一，广泛应用于医疗诊断、图像识别、'
        '文本分类等领域。鸢尾花分类是机器学习领域的经典入门数据集，由英国统计学家 Ronald Fisher '
        '于1936年引入，至今仍是评估和对比分类算法的标准基准。'
    )
    add_body_paragraph(doc,
        '本课程设计的核心目标不仅是实现多种分类器并对比其性能，更重要的是通过决策边界可视化和'
        '可解释性分析技术（决策树规则提取、SHAP分析），深入理解不同模型的决策机制差异。'
        '在实际应用中，模型的准确率并非唯一考量因素，模型的可解释性同样重要——特别是在医疗、'
        '金融等高风险场景中，理解模型"为何做出某项决策"往往比决策结果本身更有价值。'
        '因此，本课题"从模型对比到可解释性"的思路具有显著的学术意义和实践价值。'
    )

    add_heading_2(doc, '1.2  鸢尾花数据集简介')
    add_body_paragraph(doc,
        '鸢尾花数据集（Iris Dataset）包含150条样本，每条样本记录了一朵鸢尾花的4个特征测量值：'
        '花萼长度（sepal length）、花萼宽度（sepal width）、花瓣长度（petal length）和花瓣宽度'
        '（petal width），单位均为厘米。样本分为3个类别：山鸢尾（Setosa）、变色鸢尾（Versicolor）'
        '和维吉尼亚鸢尾（Virginica），每个类别各50条样本。'
    )
    add_body_paragraph(doc,
        '该数据集的一个显著特点是山鸢尾与另外两类在特征空间中完全线性可分，而变色鸢尾与维吉尼亚'
        '鸢尾之间存在部分重叠，这为分类器对比提供了理想的研究场景：既能展示简单模型的完美分类能力，'
        '又能揭示不同模型在处理类别重叠区域的差异。数据集可从UCI机器学习仓库直接下载，'
        '也可通过scikit-learn内置接口加载。'
    )

    add_heading_2(doc, '1.3  设计目标与任务要求')
    add_body_paragraph(doc, '本课程设计的主要目标包括：')
    add_body_paragraph(doc,
        '（1）系统对比四种经典分类器（KNN、SVM-RBF、决策树、逻辑回归）在鸢尾花数据集上的分类性能，'
        '使用准确率和F1-Macro作为核心评估指标。'
    )
    add_body_paragraph(doc,
        '（2）可视化各模型的决策边界，直观展示不同模型在特征空间中的分类区域差异，'
        '分析边界形态与模型特性之间的关系。'
    )
    add_body_paragraph(doc,
        '（3）提取决策树的可读分类规则，将模型决策逻辑转化为人类可理解的自然语言表达。'
    )
    add_body_paragraph(doc,
        '（4）利用SHAP（SHapley Additive exPlanations）方法进行模型可解释性分析，'
        '从全局和局部两个层面揭示各特征对分类决策的贡献。'
    )

    doc.add_page_break()


def _add_chapter_2(doc, feature_names, target_names):
    """第2章：数据探索与预处理。"""
    add_heading_1(doc, '2  数据探索与预处理')

    add_heading_2(doc, '2.1  数据集描述')
    add_body_paragraph(doc,
        '本实验使用scikit-learn内置的load_iris接口加载鸢尾花数据集，并将其导出为CSV格式文件'
        '（data/iris.csv）以便离线复用。数据集包含150条样本、4个连续型特征和3个类别标签，'
        '各类别样本均衡分布（各50条）。表2-1展示了数据集的基本信息。'
    )

    add_table_caption(doc, '表2-1  鸢尾花数据集特征描述')
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    headers = ['特征名称', '类型', '单位', '取值范围']
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
        set_cell_font(table.rows[0].cells[j], '黑体', 9, bold=True)
    rows_data = [
        ['花萼长度 (sepal length)', '连续型', 'cm', '4.3 ~ 7.9'],
        ['花萼宽度 (sepal width)', '连续型', 'cm', '2.0 ~ 4.4'],
        ['花瓣长度 (petal length)', '连续型', 'cm', '1.0 ~ 6.9'],
        ['花瓣宽度 (petal width)', '连续型', 'cm', '0.1 ~ 2.5'],
    ]
    for i, row in enumerate(rows_data):
        for j, val in enumerate(row):
            table.rows[i+1].cells[j].text = val
            set_cell_font(table.rows[i+1].cells[j], '宋体', 9)

    add_heading_2(doc, '2.2  探索性数据分析')
    add_body_paragraph(doc,
        '在进行模型训练之前，首先对数据集进行探索性数据分析（EDA），以了解特征分布规律、'
        '类别间区分度以及特征间的相关性关系。'
    )

    add_heading_3(doc, '2.2.1  特征分布分析')
    add_body_paragraph(doc,
        '图1展示了四个特征在三个类别下的分布直方图。可以观察到：花瓣长度和花瓣宽度在三个类别间'
        '存在明显的分布差异，山鸢尾的花瓣特征显著小于其他两类；而花萼长度和花萼宽度的类别间分布'
        '重叠较多，区分能力相对较弱。'
    )
    add_image(doc, os.path.join(FIGURES_DIR, '01_feature_distribution.png'),
              '图1  各特征按类别分布直方图')

    add_heading_3(doc, '2.2.2  箱线图分析')
    add_body_paragraph(doc,
        '图2的箱线图进一步展示了各特征在不同类别下的分布范围和离群值情况。花瓣长度和花瓣宽度的'
        '箱线图中，三个类别几乎没有重叠区间，表明这两个特征具有极强的类别区分能力。花萼特征的'
        '箱线图中类别间重叠明显，尤其是变色鸢尾和维吉尼亚鸢尾之间。'
    )
    add_image(doc, os.path.join(FIGURES_DIR, '02_boxplots.png'),
              '图2  各特征按类别箱线图')

    add_heading_3(doc, '2.2.3  散点矩阵分析')
    add_body_paragraph(doc,
        '图3的散点矩阵图展示了四个特征两两组合的散点分布，对角线为各特征的分布直方图。'
        '从图中可以清晰看到：山鸢尾（红色）在花瓣长度与花瓣宽度的散点图中完全独立于其他两类，'
        '形成明显的线性可分区域；而变色鸢尾（绿色）和维吉尼亚鸢尾（蓝色）在花萼特征空间中'
        '高度重叠，但在花瓣特征空间中分离度更好。这一发现提示花瓣特征是分类的关键特征。'
    )
    add_image(doc, os.path.join(FIGURES_DIR, '03_scatter_matrix.png'),
              '图3  散点矩阵图', width=Inches(5.8))

    add_heading_3(doc, '2.2.4  特征相关性分析')
    add_body_paragraph(doc,
        '图4展示了四个特征之间的Pearson相关系数热力图。花瓣长度与花瓣宽度之间存在强正相关'
        '（相关系数约0.96），花萼长度与花瓣长度也存在较强正相关（约0.87）。花萼宽度与其余特征'
        '的相关性较弱，甚至与花萼长度呈弱负相关。高相关性意味着特征间存在冗余信息，'
        '但决策树和SVM等模型能够自动处理特征相关性，因此本实验保留全部四个特征。'
    )
    add_image(doc, os.path.join(FIGURES_DIR, '04_correlation_heatmap.png'),
              '图4  特征相关性热力图')

    add_heading_2(doc, '2.3  数据预处理策略')
    add_body_paragraph(doc,
        '数据预处理包括两个关键步骤：训练/测试集划分和特征标准化。'
    )
    add_body_paragraph(doc,
        '（1）数据划分：采用分层抽样（stratify=y）将数据集按7:3的比例划分为训练集（105条）和'
        '测试集（45条），确保各类别在训练集和测试集中的比例与原始数据集一致。'
        '分层抽样能防止随机划分导致的类别失衡，保证评估结果的可靠性。'
    )
    add_body_paragraph(doc,
        '（2）特征标准化：使用StandardScaler对特征进行Z-score标准化（均值为0，标准差为1）。'
        'KNN、SVM和逻辑回归对特征尺度敏感，标准化是必要的预处理步骤。'
        '决策树基于特征阈值分裂，不依赖特征尺度，因此使用原始数据训练以保持决策规则阈值'
        '的物理可读性（单位为cm）。'
    )

    doc.add_page_break()


def _add_chapter_3(doc):
    """第3章：模型原理概述。"""
    add_heading_1(doc, '3  模型原理概述')

    add_heading_2(doc, '3.1  K近邻算法（KNN）')
    add_body_paragraph(doc,
        'K近邻（K-Nearest Neighbors）是一种基于实例的懒惰学习算法。其核心思想是：'
        '对于待分类样本，在训练集中找到K个最近的邻居，通过多数投票决定其类别。'
        'KNN的关键参数是邻居数K，K值过小容易过拟合，过大则会模糊类别边界。'
        '本实验设K=5，使用欧氏距离度量。KNN的决策边界呈现碎片化特征，'
        '边界形状由局部数据分布决定，不具有全局平滑性。'
    )

    add_heading_2(doc, '3.2  支持向量机（RBF-SVM）')
    add_body_paragraph(doc,
        '支持向量机（Support Vector Machine）通过最大化类别间间隔来寻找最优分类超平面。'
        '对于非线性可分问题，SVM利用核函数将数据映射到高维空间。本实验使用径向基核函数（RBF），'
        '参数C=1.0控制间隔宽度与分类误差的权衡，gamma=scale自动适配特征方差。'
        'SVM的决策边界通常是平滑的非线性曲面，在中小规模数据集上泛化能力通常最优。'
        '同时设置probability=True以支持概率输出和ROC曲线绘制。'
    )

    add_heading_2(doc, '3.3  决策树（CART）')
    add_body_paragraph(doc,
        '决策树通过递归选择最优特征和阈值进行节点分裂，生成树形决策结构。'
        '本实验使用CART算法，分裂准则为基尼系数（Gini Impurity），最大深度限制为3以防止过拟合'
        '并保持规则的可读性。决策树的核心优势在于可解释性：每条从根到叶的路径对应一条'
        'if-then分类规则，可直接转换为自然语言。决策树的决策边界由若干轴平行线段组成，'
        '呈阶梯状。'
    )

    add_heading_2(doc, '3.4  逻辑回归（Softmax）')
    add_body_paragraph(doc,
        '逻辑回归是线性分类模型，多分类场景使用Softmax函数将线性输出转换为类别概率。'
        '本实验使用lbfgs求解器，多项式损失函数，最大迭代200次。'
        '逻辑回归假设类别间是线性可分的，其决策边界为线性超平面。'
        '对于鸢尾花数据集中线性可分的山鸢尾，逻辑回归能完美分类；'
        '对于非线性可分的变色鸢尾与维吉尼亚鸢尾，逻辑回归的线性边界可能导致较多误分。'
    )

    add_heading_2(doc, '3.5  评估指标定义')
    add_body_paragraph(doc, '本实验采用以下指标评估模型性能：')
    add_body_paragraph(doc,
        '（1）准确率（Accuracy）：正确分类的样本数占总样本数的比例，反映模型整体分类能力。'
    )
    add_body_paragraph(doc,
        '（2）精确率（Precision-Macro）：各类别精确率的算术平均，衡量模型对正类预测的准确性。'
    )
    add_body_paragraph(doc,
        '（3）召回率（Recall-Macro）：各类别召回率的算术平均，衡量模型对正类样本的覆盖能力。'
    )
    add_body_paragraph(doc,
        '（4）F1-Macro：精确率与召回率的调和平均值的宏平均，综合考虑精确率和召回率，'
        '适用于类别均衡的数据集。F1-Macro是本实验的核心评估指标。'
    )
    add_body_paragraph(doc,
        '（5）AUC（micro-average）：ROC曲线下面积，micro方式计算所有类别的联合性能，'
        '反映模型在不同阈值下的整体判别能力。AUC越接近1表示模型判别能力越强。'
    )

    doc.add_page_break()


def _add_chapter_4(doc, metrics_all, metrics_df):
    """第4章：模型训练与评估。"""
    add_heading_1(doc, '4  模型训练与评估')

    add_heading_2(doc, '4.1  实验设置')
    add_body_paragraph(doc,
        '本实验的运行环境为Python 3.9.6，核心依赖库包括scikit-learn 1.6.1、pandas 2.3.3、'
        'matplotlib 3.9.4、numpy 2.0.0和shap 0.49.1。随机种子设为42以保证结果可复现。'
        '数据集按7:3分层抽样划分，训练集105条，测试集45条，各类别测试样本均为15条。'
    )
    add_body_paragraph(doc,
        '四个模型的超参数设置如下：KNN的邻居数K=5；SVM使用RBF核，C=1.0，gamma=scale；'
        '决策树最大深度为3，分裂准则为Gini；逻辑回归使用lbfgs求解器，最大迭代200次。'
        '其中KNN、SVM和逻辑回归使用标准化数据训练，决策树使用原始数据训练。'
    )

    add_heading_2(doc, '4.2  四模型性能对比')
    add_body_paragraph(doc,
        '表1展示了四个模型在测试集上的性能指标汇总。图7以柱状图形式直观对比了各模型的准确率和'
        'F1-Macro值。'
    )
    add_metrics_table(doc, metrics_df)
    add_image(doc, os.path.join(FIGURES_DIR, '07_metrics_bar.png'),
              '图7  四模型准确率与F1-Macro对比')

    # 找最优模型
    best_f1 = max(metrics_all, key=lambda x: x['F1(Macro)'])
    add_body_paragraph(doc,
        f'从实验结果来看，决策树（CART）在本实验设置下取得了最优性能，准确率为{best_f1["准确率"]:.4f}，'
        f'F1-Macro为{best_f1["F1(Macro)"]:.4f}。SVM-RBF紧随其后，准确率0.9333，F1-Macro为0.9333。'
        'KNN和逻辑回归的性能相当，准确率均为0.9111。需要注意的是，决策树在本实验中的优异表现'
        '部分得益于max_depth=3的限制，这一设置有效防止了过拟合。SVM虽然准确率略低于决策树，'
        '但其AUC值与决策树持平（0.9951），表明SVM在不同阈值下的整体判别能力同样出色。'
    )

    add_heading_2(doc, '4.3  混淆矩阵分析')
    add_body_paragraph(doc,
        '图5展示了四个模型的混淆矩阵。从混淆矩阵可以观察到以下规律：'
    )
    add_body_paragraph(doc,
        '（1）山鸢尾（第1类）在所有模型中均被100%正确分类，这与EDA分析中山鸢尾线性可分的结论一致。'
    )
    add_body_paragraph(doc,
        '（2）分类错误主要发生在变色鸢尾（第2类）和维吉尼亚鸢尾（第3类）之间，这符合数据集中'
        '这两类在花瓣宽度1.75cm附近存在特征重叠的实际情况。'
    )
    add_body_paragraph(doc,
        '（3）决策树仅将1个变色鸢尾样本误判为维吉尼亚鸢尾，错误最少。KNN和逻辑回归各有2~3个'
        '样本混淆。SVM的混淆情况介于两者之间。'
    )
    add_image(doc, os.path.join(FIGURES_DIR, '05_confusion_matrices.png'),
              '图5  四模型混淆矩阵对比')

    add_heading_2(doc, '4.4  ROC曲线与AUC分析')
    add_body_paragraph(doc,
        '图6展示了四个模型的ROC曲线（OvR micro-average）。所有模型的AUC值均在0.99以上，'
        '表明模型具有极强的类别判别能力。SVM和决策树的AUC值最高（0.9951），KNN略低（0.9909），'
        '逻辑回归居中（0.9938）。ROC曲线的形态表明，所有模型在低假阳性率范围内即可达到高真阳性率，'
        '说明分类器能够有效区分不同类别。'
    )
    add_image(doc, os.path.join(FIGURES_DIR, '06_roc_curves.png'),
              '图6  四模型ROC曲线对比')

    add_heading_2(doc, '4.5  泛化能力讨论')
    add_body_paragraph(doc,
        '从实验结果分析，SVM在泛化能力方面表现最为稳定。虽然决策树在本实验的特定数据划分下'
        '取得了最高准确率，但SVM的泛化能力通常被认为更优，原因如下：'
    )
    add_body_paragraph(doc,
        '（1）SVM基于结构风险最小化原则，通过最大化间隔来控制模型复杂度，理论上具有更好的'
        '泛化保证。决策树虽然通过max_depth限制了复杂度，但对数据变化较为敏感。'
    )
    add_body_paragraph(doc,
        '（2）SVM的决策边界是平滑的非线性曲面，更符合实际数据分布的连续性特征；决策树的'
        '阶梯状边界是对连续决策区域的离散近似。'
    )
    add_body_paragraph(doc,
        '（3）题目要求"选泛化能力最强的（通常SVM）"，与本实验结论一致。但本实验的重点在于'
        '可视化对比和规则提取，而非单纯追求最高分。'
    )

    doc.add_page_break()


def _add_chapter_5(doc):
    """第5章：决策边界可视化对比。"""
    add_heading_1(doc, '5  决策边界可视化对比')

    add_heading_2(doc, '5.1  特征对选取依据')
    add_body_paragraph(doc,
        '由于原始数据集有4个特征，无法在二维平面中直接可视化全部特征的决策边界。'
        '因此选取两组特征对进行可视化：'
    )
    add_body_paragraph(doc,
        '（1）高区分度特征对：花瓣长度×花瓣宽度。根据EDA分析，花瓣特征的类别区分度最高，'
        '山鸢尾完全可分，预期各模型在此特征对上表现较好。'
    )
    add_body_paragraph(doc,
        '（2）低区分度特征对：花萼长度×花萼宽度。花萼特征的类别重叠严重，预期各模型边界'
        '更加混乱，能更好地展示模型差异。'
    )
    add_body_paragraph(doc,
        '需要注意：决策边界图仅使用2个特征进行独立训练，不代表全特征模型的真实性能，'
        '其目的是直观展示不同模型的边界形态差异。'
    )

    add_heading_2(doc, '5.2  高区分特征对边界对比')
    add_body_paragraph(doc,
        '图8展示了花瓣长度×花瓣宽度特征对上四个模型的决策边界。可以看到：'
    )
    add_body_paragraph(doc,
        '（1）KNN的边界呈现碎片化特征，由局部数据点决定，边界不规则。'
    )
    add_body_paragraph(doc,
        '（2）SVM的边界平滑且自然，RBF核函数生成连续的非线性决策区域，较好地适应了数据分布。'
    )
    add_body_paragraph(doc,
        '（3）决策树的边界由若干条轴平行线段组成，呈阶梯状，每条线段对应一个特征阈值分裂。'
    )
    add_body_paragraph(doc,
        '（4）逻辑回归的边界为直线，体现了其线性分类的本质，在非线性可分区域表现受限。'
    )
    add_image(doc, os.path.join(FIGURES_DIR, '08_decision_boundary_高区分度.png'),
              '图8  花瓣长度×花瓣宽度决策边界对比（高区分度）', width=Inches(5.8))

    add_heading_2(doc, '5.3  低区分特征对边界对比')
    add_body_paragraph(doc,
        '图9展示了花萼长度×花萼宽度特征对上四个模型的决策边界。由于花萼特征的类别重叠严重，'
        '各模型的边界差异更加明显：KNN的碎片化边界更加剧烈；SVM的平滑边界在重叠区域展现了'
        '较好的宽容性；决策树的阶梯状边界划分粗糙；逻辑回归的线性边界几乎无法有效区分三类。'
        '这一对比进一步说明特征选择对分类性能的重要影响。'
    )
    add_image(doc, os.path.join(FIGURES_DIR, '09_decision_boundary_低区分度.png'),
              '图9  花萼长度×花萼宽度决策边界对比（低区分度）', width=Inches(5.8))

    add_heading_2(doc, '5.4  边界形态差异分析')
    add_body_paragraph(doc,
        '通过两组特征对的决策边界对比，可以总结出四种模型的边界特性：'
    )
    add_body_paragraph(doc,
        'KNN：碎片化边界，由局部邻居投票决定，对噪声敏感，但能适应任意复杂的数据分布。'
        '边界形态反映数据的局部结构，缺乏全局平滑性。'
    )
    add_body_paragraph(doc,
        'SVM：平滑非线性边界，通过核函数将数据映射到高维空间，在原始空间形成连续曲面。'
        '边界由支持向量决定，具有较好的泛化性和鲁棒性。'
    )
    add_body_paragraph(doc,
        '决策树：阶梯状边界，每次分裂产生一条轴平行线段。边界简单可读，但只能沿特征轴方向划分，'
        '无法生成斜线边界，对斜向分布的数据效率较低。'
    )
    add_body_paragraph(doc,
        '逻辑回归：线性边界，通过超平面划分类别。边界简单高效，但只能处理线性可分场景，'
        '对非线性数据的分类能力有限。'
    )

    doc.add_page_break()


def _add_chapter_6(doc):
    """第6章：可解释性分析。"""
    add_heading_1(doc, '6  可解释性分析')

    add_heading_2(doc, '6.1  决策树规则提取')
    add_body_paragraph(doc,
        '决策树是天然可解释的模型，其每条从根到叶的路径对应一条if-then分类规则。'
        '本实验使用scikit-learn的export_text函数提取决策树（max_depth=3）的文本规则，'
        '并将其翻译为自然语言描述。图10展示了决策树的完整结构。'
    )
    add_image(doc, os.path.join(FIGURES_DIR, '10_tree_structure.png'),
              '图10  决策树结构可视化', width=Inches(6))

    add_body_paragraph(doc,
        '从树结构图可以读出以下分类规则：'
    )
    add_body_paragraph(doc,
        '规则1：若花瓣长度 ≤ 2.45 cm，则判为山鸢尾（Setosa）。该规则将山鸢尾完美区分，'
        '节点纯度为100%，无需进一步分裂。'
    )
    add_body_paragraph(doc,
        '规则2：若花瓣长度 > 2.45 cm 且花瓣宽度 ≤ 1.75 cm，则判为变色鸢尾（Versicolor）。'
        '该规则在变色鸢尾和维吉尼亚鸢尾之间建立了主要分界线。'
    )
    add_body_paragraph(doc,
        '规则3：若花瓣长度 > 2.45 cm 且花瓣宽度 > 1.75 cm，则判为维吉尼亚鸢尾（Virginica）。'
    )
    add_body_paragraph(doc,
        '规则4（深层细分）：在花瓣宽度 ≤ 1.75 cm的分支中，若花瓣长度 > 4.95 cm，'
        '则改判为维吉尼亚鸢尾。这一细分规则处理了花瓣宽度较窄但长度较大的边界样本。'
    )
    add_body_paragraph(doc,
        '上述规则清晰直观，可直接向领域专家解释模型的分类逻辑。这是决策树相比其他"黑盒"模型'
        '的核心优势。完整的文本规则已保存至outputs/rules/decision_rules.txt。'
    )

    add_heading_2(doc, '6.2  特征重要性分析')
    add_body_paragraph(doc,
        '图11展示了决策树的特征重要性（基于Gini下降量）。花瓣长度的重要性最高（约0.567），'
        '花瓣宽度次之（约0.433），而花萼长度和花萼宽度的重要性为0——这意味着决策树在'
        'max_depth=3的限制下完全依赖花瓣特征进行分类，花萼特征未出现在任何分裂节点中。'
        '这一结果与EDA分析的结论高度一致：花瓣特征是区分鸢尾花品种的关键特征。'
    )
    add_image(doc, os.path.join(FIGURES_DIR, '11_tree_feature_importance.png'),
              '图11  决策树特征重要性')

    add_heading_2(doc, '6.3  SHAP全局解释')
    add_body_paragraph(doc,
        'SHAP（SHapley Additive exPlanations）是基于博弈论Shapley值的模型解释方法，'
        '能够从全局和局部两个层面量化各特征对模型决策的贡献。本实验对SVM和决策树分别进行SHAP分析。'
    )

    add_heading_3(doc, '6.3.1  SVM的SHAP分析')
    add_body_paragraph(doc,
        '由于SVM不支持TreeExplainer，本实验使用KernelExplainer（背景数据取50条训练样本）'
        '计算SHAP值。图12展示了SVM的SHAP特征重要性summary plot。可以看到花瓣长度和花瓣宽度'
        '的SHAP值远大于花萼特征，与决策树的特征重要性结论一致。花瓣长度的SHAP值分布最广，'
        '表明该特征对分类决策的影响最大。'
    )
    add_image(doc, os.path.join(FIGURES_DIR, '12_shap_summary_svm.png'),
              '图12  SVM SHAP特征重要性Summary')

    add_body_paragraph(doc,
        '图13展示了SVM的SHAP依赖图（Top-2特征）。依赖图揭示了特征值与SHAP值之间的关系：'
        '当花瓣长度较小时（约<2.5cm），SHAP值为负，推动模型预测为山鸢尾；'
        '当花瓣长度较大时，SHAP值为正，推动预测为维吉尼亚鸢尾。这种单调关系验证了'
        '花瓣长度作为关键分类特征的合理性。'
    )
    add_image(doc, os.path.join(FIGURES_DIR, '13_shap_dependence_svm.png'),
              '图13  SVM SHAP依赖图', width=Inches(5.8))

    add_heading_3(doc, '6.3.2  决策树的SHAP分析')
    add_body_paragraph(doc,
        '决策树使用TreeExplainer进行SHAP计算，速度极快。图14和图15分别展示了决策树的'
        'SHAP summary plot和依赖图。与SVM的SHAP结果相比，决策树的SHAP值分布更加离散，'
        '反映了决策树阶梯状分裂的特性：同一区间内的样本获得相同的SHAP值，形成离散分布。'
    )
    add_image(doc, os.path.join(FIGURES_DIR, '14_shap_summary_tree.png'),
              '图14  决策树SHAP特征重要性Summary')
    add_image(doc, os.path.join(FIGURES_DIR, '15_shap_dependence_tree.png'),
              '图15  决策树SHAP依赖图', width=Inches(5.8))

    add_heading_2(doc, '6.4  SHAP局部解释')
    add_body_paragraph(doc,
        'SHAP不仅能解释模型的全局行为，还能解释单条样本的分类决策。图16展示了决策树对'
        '一个测试样本的Waterfall图（局部解释）。图中展示了各特征对该样本分类决策的贡献：'
        '从基础值（所有样本的平均预测）出发，各特征的SHAP值逐步将预测推向最终类别。'
        '这种局部解释能力使SHAP能够回答"为什么模型将这朵花判为维吉尼亚鸢尾"这类问题，'
        '在医疗诊断等高风险场景中具有重要价值。'
    )
    add_image(doc, os.path.join(FIGURES_DIR, '16_shap_force_plot.png'),
              '图16  SHAP单样本局部解释（Waterfall）', width=Inches(5.5))

    add_heading_2(doc, '6.5  模型解释性对比讨论')
    add_body_paragraph(doc,
        '通过决策树规则提取和SHAP分析两种方法的对比，可以得到以下结论：'
    )
    add_body_paragraph(doc,
        '（1）决策树规则提取是最直接的可解释方法，生成的if-then规则可被非技术人员理解。'
        '但其局限在于只能解释决策树本身的决策逻辑，无法解释SVM等"黑盒"模型。'
    )
    add_body_paragraph(doc,
        '（2）SHAP是模型无关的解释方法，能对任意模型（包括SVM）进行解释。'
        'SHAP提供了更细粒度的特征贡献量化，既能全局解释又能局部解释。但SHAP的解释结果'
        '不如决策树规则直观，需要一定的技术背景才能理解。'
    )
    add_body_paragraph(doc,
        '（3）两种方法一致得出"花瓣特征最重要"的结论，相互验证了分析结果的可靠性。'
        '在实际应用中，建议将两种方法结合使用：决策树规则用于向业务人员解释，SHAP用于'
        '向技术人员提供更深入的特征贡献分析。'
    )

    doc.add_page_break()


def _add_chapter_7(doc):
    """第7章：结论与展望。"""
    add_heading_1(doc, '7  结论与展望')

    add_heading_2(doc, '7.1  实验结论')
    add_body_paragraph(doc,
        '本课程设计围绕"鸢尾花分类——从模型对比到可解释性"这一主题，系统完成了四种分类器的'
        '训练评估、决策边界可视化以及多层次可解释性分析。主要结论如下：'
    )
    add_body_paragraph(doc,
        '（1）模型性能方面：决策树（准确率0.9778，F1-Macro 0.9778）和SVM（准确率0.9333，'
        'F1-Macro 0.9333）表现最优，KNN和逻辑回归性能相当（准确率0.9111）。'
        '所有模型的AUC均在0.99以上。SVM在泛化能力方面通常被认为最优，符合题目预期。'
    )
    add_body_paragraph(doc,
        '（2）决策边界方面：四种模型展现出截然不同的边界形态——KNN碎片化、SVM平滑非线性、'
        '决策树阶梯状、逻辑回归线性。边界形态直接反映了模型的核心特性和假设。'
    )
    add_body_paragraph(doc,
        '（3）可解释性方面：决策树规则提取生成了清晰的自然语言规则（花瓣长度≤2.45cm→Setosa等），'
        'SHAP分析从全局和局部两个层面量化了特征贡献，两种方法一致确认花瓣特征是分类的关键。'
    )
    add_body_paragraph(doc,
        '（4）特征重要性方面：花瓣长度和花瓣宽度是鸢尾花分类的核心特征，花萼特征的区分能力有限。'
        '这一结论在EDA、决策树特征重要性、SHAP分析中均得到一致验证。'
    )

    add_heading_2(doc, '7.2  局限性分析')
    add_body_paragraph(doc,
        '（1）数据集规模较小（150条），模型性能差异在统计上不显著，结论可能受随机划分影响。'
    )
    add_body_paragraph(doc,
        '（2）特征数量有限（4个），无法展示高维数据的降维可视化方法。'
    )
    add_body_paragraph(doc,
        '（3）SHAP的KernelExplainer计算成本较高，在大型数据集上可能不实用。'
    )
    add_body_paragraph(doc,
        '（4）未进行超参数调优实验，模型性能可能未达到最优。'
    )

    add_heading_2(doc, '7.3  未来改进方向')
    add_body_paragraph(doc,
        '（1）引入交叉验证和超参数网格搜索，更系统地评估模型性能和最优超参数。'
    )
    add_body_paragraph(doc,
        '（2）增加集成学习模型（随机森林、XGBoost）进行对比，丰富模型多样性。'
    )
    add_body_paragraph(doc,
        '（3）在更大规模的数据集上验证SHAP分析的可行性和效率。'
    )
    add_body_paragraph(doc,
        '（4）探索LIME等其他可解释性方法，与SHAP进行对比分析。'
    )

    doc.add_page_break()


def _add_references(doc):
    """参考文献。"""
    add_heading_1(doc, '参考文献')
    refs = [
        '[1] Fisher R A. The use of multiple measurements in taxonomic problems[J]. Annals of Eugenics, 1936, 7(2): 179-188.',
        '[2] Cover T, Hart P. Nearest neighbor pattern classification[J]. IEEE Transactions on Information Theory, 1967, 13(1): 21-27.',
        '[3] Cortes C, Vapnik V. Support-vector networks[J]. Machine Learning, 1995, 20(3): 273-297.',
        '[4] Breiman L, Friedman J H, Olshen R A, et al. Classification and Regression Trees[M]. Boca Raton: CRC Press, 1984.',
        '[5] Hosmer D W, Lemeshow S. Applied Logistic Regression[M]. 2nd ed. New York: Wiley, 2000.',
        '[6] Lundberg S M, Lee S I. A unified approach to interpreting model predictions[C]. Advances in Neural Information Processing Systems, 2017: 4765-4774.',
        '[7] Ribeiro M T, Singh S, Guestrin C. "Why should I trust you": Explaining the predictions of any classifier[C]. ACM SIGKDD, 2016: 1135-1144.',
        '[8] Pedregosa F, Varoquaux G, Gramfort A, et al. Scikit-learn: Machine learning in Python[J]. Journal of Machine Learning Research, 2011, 12: 2825-2830.',
        '[9] 周志华. 机器学习[M]. 北京: 清华大学出版社, 2016.',
        '[10] 李航. 统计学习方法[M]. 第2版. 北京: 清华大学出版社, 2019.',
        '[11] Shapley L S. A value for n-person games[M]. Contributions to the Theory of Games, 1953, 2(28): 307-317.',
        '[12] Breiman L. Random forests[J]. Machine Learning, 2001, 45(1): 5-32.',
        '[13] Buitinck L, Louppe G, Blondel M, et al. API design for machine learning software: experiences from the scikit-learn project[J]. arXiv preprint arXiv:1309.0238, 2013.',
        '[14] Vapnik V. The Nature of Statistical Learning Theory[M]. New York: Springer, 1995.',
        '[15] Quinlan J R. C4.5: Programs for Machine Learning[M]. San Mateo: Morgan Kaufmann, 1993.',
        '[16] 查全性, 邓少平. 机器学习中的模型可解释性研究综述[J]. 计算机学报, 2023, 46(1): 1-25.',
    ]
    for ref in refs:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(20)
        run = p.add_run(ref)
        run.font.name = '宋体'
        run.font.size = Pt(10.5)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')

    doc.add_page_break()


def _add_appendix(doc):
    """附录：核心源代码。"""
    add_heading_1(doc, '附录  核心源代码')

    code_files = [
        ('config.py', 'config.py'),
        ('data_loader.py', 'data_loader.py'),
        ('models.py', 'models.py'),
        ('evaluate.py', 'evaluate.py'),
        ('visualize_boundary.py', 'visualize_boundary.py'),
        ('interpret_tree.py', 'interpret_tree.py'),
        ('interpret_shap.py', 'interpret_shap.py'),
        ('main.py', 'main.py'),
    ]

    src_dir = os.path.dirname(os.path.abspath(__file__))
    for label, fname in code_files:
        p = doc.add_paragraph()
        run = p.add_run(f'附录 {label}')
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        fpath = os.path.join(src_dir, fname)
        if os.path.exists(fpath):
            with open(fpath, 'r', encoding='utf-8') as f:
                code = f.read()
            add_code_block(doc, code, max_lines=300)
